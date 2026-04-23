import json
from base64 import b64decode

from docx import Document as WordDocument
from docx.oxml.shared import OxmlElement, qn

from documind.pipeline.convert import ConvertPipeline

from tests.conftest import SAMPLE_DOCX
from tests.fixtures import build_complex_table_docx


def test_convert_pipeline_writes_chunks_assets_and_indexes(tmp_path):
    pipeline = ConvertPipeline.build_default()
    input_path = SAMPLE_DOCX

    pipeline.run(input_path=input_path, out_dir=tmp_path, split_level=2)

    index_file = tmp_path / "index.md"
    summary_file = tmp_path / "summary.json"
    chunk_files = sorted(path for path in tmp_path.rglob("*.md") if path != index_file)

    assert index_file.exists()
    assert summary_file.exists()
    assert (tmp_path / "assets").is_dir()
    assert len(chunk_files) >= 4
    assert any("![" in path.read_text(encoding="utf-8") for path in chunk_files)

    summary = json.loads(summary_file.read_text(encoding="utf-8"))
    assert summary["source_doc"] == "sample.docx"
    assert summary["chunks"]


def test_convert_pipeline_handles_public_example_with_multiple_chunks_and_assets(tmp_path):
    pipeline = ConvertPipeline.build_default()
    input_path = SAMPLE_DOCX

    pipeline.run(input_path=input_path, out_dir=tmp_path, split_level=2)

    chunk_files = [path for path in tmp_path.rglob("*.md") if path.name != "index.md" or path.parent != tmp_path]
    asset_files = list((tmp_path / "assets").glob("*"))

    assert len(chunk_files) >= 4
    assert len(asset_files) >= 2
    assert any(path.name == "index.md" and path.parent != tmp_path for path in chunk_files)
    features_chunk = (tmp_path / "features" / "index.md").read_text(encoding="utf-8")
    gallery_chunk = (tmp_path / "features" / "media-gallery" / "index.md").read_text(encoding="utf-8")
    summary = json.loads((tmp_path / "summary.json").read_text(encoding="utf-8"))
    features_summary = next(
        chunk for chunk in summary["chunks"] if chunk["path"] == "features/index.md"
    )

    assert "https://example.com/specification" in features_chunk
    assert "![sample-diagram]" in gallery_chunk
    assert features_summary["links"] == [
        {"text": "project specification", "url": "https://example.com/specification", "kind": "external"}
    ]


def test_convert_pipeline_renders_complex_large_tables_as_html(tmp_path):
    pipeline = ConvertPipeline.build_default()
    input_path = build_complex_table_docx(tmp_path / "complex-table.docx")

    pipeline.run(input_path=input_path, out_dir=tmp_path, split_level=2)

    chunk_texts = [
        path.read_text(encoding="utf-8")
        for path in tmp_path.rglob("*.md")
        if path.name != "index.md" or path.parent != tmp_path
    ]

    assert any("<table>" in text for text in chunk_texts)


def test_convert_pipeline_keeps_images_inside_table_cells(tmp_path):
    image_bytes = b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wn8lY4AAAAASUVORK5CYII="
    )
    image_path = tmp_path / "pixel.png"
    image_path.write_bytes(image_bytes)

    doc_path = tmp_path / "table-image.docx"
    document = WordDocument()
    table = document.add_table(rows=1, cols=1)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.add_run("Cell text ")
    paragraph.add_run().add_picture(str(image_path))
    document.save(doc_path)

    pipeline = ConvertPipeline.build_default()
    out_dir = tmp_path / "out"
    pipeline.run(input_path=doc_path, out_dir=out_dir, split_level=2)

    chunk_text = (out_dir / "00_preface.md").read_text(encoding="utf-8")

    assert "<img src=\"./assets/img_001.png\"" in chunk_text


def test_convert_pipeline_renders_table_links_as_html_and_collects_summary_links(tmp_path):
    doc_path = tmp_path / "table-link.docx"
    document = WordDocument()
    table = document.add_table(rows=1, cols=1)
    paragraph = table.cell(0, 0).paragraphs[0]

    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "HYPERLINK https://example.com/spec.xlsx \\tdfn %u89C4%u683C.xlsx "
    instruction_run._r.append(instruction)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    paragraph.add_run("规格")
    paragraph.add_run(".xlsx")

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)

    document.save(doc_path)

    pipeline = ConvertPipeline.build_default()
    out_dir = tmp_path / "out"
    pipeline.run(input_path=doc_path, out_dir=out_dir, split_level=2)

    chunk_text = (out_dir / "00_preface.md").read_text(encoding="utf-8")
    summary = json.loads((out_dir / "summary.json").read_text(encoding="utf-8"))

    assert "<table>" in chunk_text
    assert '<a href="https://example.com/spec.xlsx">规格.xlsx</a>' in chunk_text
    assert summary["chunks"][0]["links"] == [
        {
            "text": "规格.xlsx",
            "url": "https://example.com/spec.xlsx",
            "kind": "attachment",
        }
    ]


def test_convert_pipeline_writes_nested_index_paths_for_h1_and_h2_chunks(tmp_path):
    pipeline = ConvertPipeline.build_default()
    input_path = SAMPLE_DOCX

    pipeline.run(input_path=input_path, out_dir=tmp_path, split_level=2)

    assert (tmp_path / "overview" / "index.md").exists()
    assert (tmp_path / "overview" / "glossary" / "index.md").exists()
    assert "![sample-diagram](../../assets/img_001.png)" in (
        tmp_path / "features" / "media-gallery" / "index.md"
    ).read_text(encoding="utf-8")

    summary = json.loads((tmp_path / "summary.json").read_text(encoding="utf-8"))
    paths = {chunk["path"] for chunk in summary["chunks"]}
    assert "overview/index.md" in paths
    assert "overview/glossary/index.md" in paths


def test_convert_pipeline_writes_single_inline_markdown_file(tmp_path):
    pipeline = ConvertPipeline.build_default()
    input_path = SAMPLE_DOCX

    pipeline.run(input_path=input_path, out_dir=tmp_path, split_level=2, mode="inline")

    index_text = (tmp_path / "index.md").read_text(encoding="utf-8")

    assert (tmp_path / "index.md").exists()
    assert not (tmp_path / "summary.json").exists()
    assert not (tmp_path / "overview" / "index.md").exists()
    assert "# Overview" in index_text
    assert "## Glossary" in index_text
    assert "![sample-diagram](./assets/img_001.png)" in index_text
    assert not index_text.startswith("---\n")


def test_convert_pipeline_keeps_complex_tables_in_inline_mode(tmp_path):
    pipeline = ConvertPipeline.build_default()
    input_path = build_complex_table_docx(tmp_path / "complex-table.docx")

    pipeline.run(input_path=input_path, out_dir=tmp_path, split_level=2, mode="inline")

    index_text = (tmp_path / "index.md").read_text(encoding="utf-8")

    assert "<table>" in index_text
