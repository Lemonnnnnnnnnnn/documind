from __future__ import annotations

from base64 import b64decode
from pathlib import Path

from docx import Document as WordDocument


PNG_RED = b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8z8AABYMBgVyuk3sAAAAASUVORK5CYII="
)
PNG_GREEN = b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wn8lY4AAAAASUVORK5CYII="
)
PNG_BLUE = b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAwMBAS8p0uoAAAAASUVORK5CYII="
)


def write_png(path: Path, payload: bytes = PNG_GREEN) -> Path:
    path.write_bytes(payload)
    return path


def build_complex_table_docx(path: Path) -> Path:
    document = WordDocument()
    document.add_paragraph("Synthetic complex table fixture")
    table = document.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Region"
    table.cell(0, 1).text = "Tier A"
    table.cell(0, 2).text = "Tier B"
    table.cell(1, 0).merge(table.cell(2, 0)).text = "North"
    table.cell(1, 1).text = "Enabled"
    table.cell(1, 2).text = "Review"
    table.cell(2, 1).text = "Delayed"
    table.cell(2, 2).text = "Blocked"
    document.save(path)
    return path
